import os
import io
import fitz  # PyMuPDF
import docx
import tempfile
import pythoncom
import win32com.client

def parse_pdf(file_bytes: bytes) -> str:
    """Extracts raw text from PDF bytes page by page."""
    try:
        # Open PDF from bytes stream
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text_parts = []
        for i, page in enumerate(doc):
            page_text = page.get_text()
            if page_text:
                text_parts.append(f"--- PAGE {i + 1} ---\n{page_text}")
        return "\n\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Error parsing PDF document: {str(e)}")

def parse_docx(file_bytes: bytes) -> str:
    """Extracts raw text from DOCX bytes including paragraphs and table contents."""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        text_parts = []
        
        # Extract paragraph text
        for p in doc.paragraphs:
            if p.text.strip():
                text_parts.append(p.text)
                
        # Extract table text
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    text_parts.append(" | ".join(row_cells))
                    
        return "\n\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Error parsing DOCX document: {str(e)}")

def parse_doc_with_word(file_bytes: bytes) -> str:
    """Uses MS Word COM Automation to parse binary .doc files by converting them to text."""
    pythoncom.CoInitialize()
    
    # Use D:\temp since C: drive has extremely limited space
    temp_dir = "D:\\temp"
    os.makedirs(temp_dir, exist_ok=True)
    
    fd, temp_path = tempfile.mkstemp(suffix=".doc", dir=temp_dir)
    try:
        with os.fdopen(fd, "wb") as f:
            f.write(file_bytes)
            
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        
        doc = word.Documents.Open(temp_path)
        text = doc.Content.Text
        doc.Close(False)
        word.Quit()
        
        if text:
            # Normalize Word carriage returns
            text = text.replace('\r', '\n')
        return text
    except Exception as e:
        raise ValueError(f"Error parsing .doc document via MS Word: {str(e)}")
    finally:
        if os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except Exception:
                pass
        pythoncom.CoUninitialize()

def parse_document(filename: str, file_bytes: bytes) -> str:
    """Orchestrates parsing based on file extension."""
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        return parse_pdf(file_bytes)
    elif ext == ".docx":
        return parse_docx(file_bytes)
    elif ext == ".doc":
        return parse_doc_with_word(file_bytes)
    elif ext == ".txt":
        return file_bytes.decode("utf-8", errors="ignore")
    else:
        raise ValueError(f"Unsupported file format: {ext}. Only PDF, DOCX, DOC, and TXT are supported.")

